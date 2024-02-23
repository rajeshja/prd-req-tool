import streamlit as st
import docx 
import os
from langchain.document_loaders.word_document import Docx2txtLoader
from langchain_core.documents.base import Document
from langchain.globals import set_debug
import  tiktoken

def calculate_tokens(text):
    return len(tiktoken.encoding_for_model('gpt-3.5-turbo').encode(text))

def load_document(prd_file):
    prd_bytes = prd_file.getvalue()
    os.makedirs("temp", exist_ok=True)
    temp_file_name = os.path.join("temp", prd_file.name)
    with open(temp_file_name, 'wb') as f:
        f.write(prd_bytes)

    document = docx.Document(temp_file_name)
    os.remove(temp_file_name)
    return temp_file_name,document

def calculate_tokens_for_document(document):
    all_text = ""
    for para in document.paragraphs:
        all_text += para.text
        all_text += "\n"

    return calculate_tokens(all_text)

def chunk_using_nltk(chunks):
    import nltk
    nltk.download('punkt')

    from langchain.text_splitter import NLTKTextSplitter

    text_splitter = NLTKTextSplitter(chunk_size=2000)
    chunks = text_splitter.split_documents(chunks)
    return chunks
    
def parse_word_doc_single_chunk(temp_file_name, document):
    text = ""
    metadata = {"source": os.path.basename(temp_file_name)}
    for para in document.paragraphs:
        text += para.text
        text += "\n\n"
    return [Document(page_content=text, metadata=metadata)]
    # chunks.append(Document(page_content=text, metadata=metadata))
    # return chunks

def parse_word_doc_section_chunks(temp_file_name, document):
    chunks = []
    current_fragment = ""
    current_metadata = {"source": os.path.basename(temp_file_name)}
    prev_para_type = ""

    for para in document.paragraphs:
        if para.style.name.startswith("Heading") and prev_para_type != "heading":
            # is header 1
            if current_fragment != "":
                chunk = Document(page_content=current_fragment, metadata=current_metadata)
                chunks.append(chunk)
            current_fragment = para.text 
            current_metadata = {"source": os.path.basename(temp_file_name), "heading": para.text}
            prev_para_type = "heading"
        else:
            current_fragment += f"\n{para.text}"
            prev_para_type = "para"

    if current_fragment != "":
        chunk = Document(page_content=current_fragment, metadata=current_metadata)
        chunks.append(chunk)

    return chunks

def show_chunks(calculate_tokens, chunks):
    for chunk in chunks:
        st.write(f"Bytes: {len(chunk.page_content)}, Tokens: {calculate_tokens(chunk.page_content)}")
        st.subheader("Metadata")
        st.json(chunk.metadata)
        st.subheader("Document content")
        st.write(chunk.page_content)
        st.divider()

st.header("Time to import the PRD")

prd_file = st.file_uploader("Upload a PRD document", ["pdf", "docx"])

def documents_to_chroma(chunks, embeddings):
    from langchain_community.vectorstores.chroma import Chroma
    ids = [str(i+1) for i in range(len(chunks))]
    vectorstore = Chroma.from_documents(documents=chunks, embedding=embeddings, ids=ids)
    return vectorstore

def get_gpt35_llm():
    return get_gpt35turbo()

def get_gpt35turbo():
    """This returns the ChatGPT 3.5 Turbo Instruct model. The context window is 4k.
    This model is an Instruct model, and is best suited to for NLP tasks and when
    you want it to stick to your desired output.
    """
    from langchain_openai import ChatOpenAI
    return ChatOpenAI(model="gpt-3.5-turbo", temperature=0)

def get_gpt35_instruct():
    """This returns the ChatGPT 3.5 Turbo Instruct model. The context window is 4k.
    This model is an Instruct model, and is best suited to for NLP tasks and when
    you want it to stick to your desired output.
    """
    from langchain_openai import ChatOpenAI
    return ChatOpenAI(model="gpt-3.5-turbo-instruct", temperature=0)

def get_gpt35_16k():
    """This returns the ChatGPT 3.5 Turbo model. The context window is 16k.
    However, this model is not an Instruct model, and is best suited to Chat mode.
    """
    from langchain_openai import ChatOpenAI
    return ChatOpenAI(model="gpt-3.5-turbo-0125", temperature=0)

def get_gpt4turbo_128k():
    """This returns the ChatGPT 4 Turbo model. The context window is 128k.
    The model is currently in preview, and is cheaper than standard gpt4 
    """
    from langchain_openai import ChatOpenAI
    return ChatOpenAI(model="gpt-4-0125-preview", temperature=0)

def dummy_chains_code():
    from langchain_openai import OpenAIEmbeddings
    from langchain import hub
    from langchain.chains import RetrievalQA, MapReduceDocumentsChain, ReduceDocumentsChain, StuffDocumentsChain

    map_template = """
        The following text """

    func_req_query = """Look at the given context, and see if this text contains any functional requirements.
                        If there are any functional requirements mentioned, give a short summary of the requirement.
                        Also provide a short title of the requirement. the title should be within 4 words.
                        Do not include non-functional requirements in your response. 
                        If there are no functional requirements, use the title 'none'.
                        Provide the answer in the following structure: \{ 'title': 'Requirement Title', 'summary': 'Requirement Summary' \}"""

    from langchain_core.prompts import PromptTemplate, HumanMessagePromptTemplate

    prompt_template = hub.pull("rlm/rag-prompt")
    print("From hub:")
    print(prompt_template)

    # """
    #     input_variables=['context', 'question'] 
    #     """
    prompt_template = PromptTemplate(input_variables=["context", "question"], 
                                     template="""You are the assistant to a Product Manager for question-answering tasks.
                                                 Use the following pieces of retrieved context to answer the question.
                                                 If you don't know the answer, just say that you don't know.
                                                 Use three sentences maximum and keep the answer concise.
                                                 Question: {question} 
                                                 Context: {context} 
                                                 Answer:""")
    
    print("Custom prompt:")
    print(prompt_template)

    # chain = RetrievalQA.from_chain_type(
    #     llm=llm,
    #     retriever = vectorstore.as_retriever(),
    #     return_source_documents=False,
    #     chain_type_kwargs={"prompt": prompt_template}
    # )

    embedding_vector = OpenAIEmbeddings().embed_query(func_req_query)
    docs_retrieved = vectorstore.similarity_search_by_vector_with_relevance_scores(embedding_vector, k=50)

    st.header("Docs from Chroma")

    # docs_retrieved = retriever.invoke(func_req_query)

    for i, (doc, relevance) in enumerate(docs_retrieved):
        st.write(f"Sr No: {i+1}, Relevance: {relevance}, Tokens: {calculate_tokens(doc.page_content)}")
        st.write(doc)
        st.divider()

    from string import Template

    prompt = Template("""You are the assistant to a Product Manager for question-answering tasks.
                                                 Use the following pieces of retrieved context to answer the question.
                                                 If you don't know the answer, just say that you don't know.
                                                 Use three sentences maximum and keep the answer concise.
                                                 Question: $question
                                                 Context: $context
                                                 Answer:""")

    for i in range(2, 4):
        # question = func_req_query
        # context = docs_retrieved[i]
        # response = chain.invoke({"query": func_req_query, "context": docs_retrieved[i]})
        st.divider()
        st.write(llm(prompt.substitute(question=func_req_query, context=docs_retrieved[i])))
        # st.write(response["result"])
        # print(response["result"])
        st.divider()


def get_requirements_from_stuff_chain(llm, chunks):
    from langchain.chains.combine_documents.stuff import StuffDocumentsChain
    from langchain.chains.llm import LLMChain
    from langchain.prompts import PromptTemplate

    prompt_template = """The following is a product requirements document. Go through it and 
                         summarize it as a list of product requirement summaries. Use a single 
                         sentence for each requirement summary. Provide the response in a JSON
                         as a list of strings:
                         ====== begin product requirements document ======
                         {document_text}
                         ====== end product requirements document ======
                         List of requirements:
                         """

    prompt = PromptTemplate.from_template(prompt_template)
    llm_chain = LLMChain(llm=llm, prompt=prompt)
    stuff_chain = StuffDocumentsChain(llm_chain=llm_chain, document_variable_name="document_text")
    st.header("Requirement Summaries from LLM")
    requirements = stuff_chain.run(chunks)
    import json
    return json.loads(requirements)["requirements"]


from typing import List, Optional
from langchain_core.pydantic_v1 import BaseModel, Field

class Requirements(BaseModel):
    requirements: Optional[List[str]] = None

@st.cache_data
def get_requirements_from_mapreduce_chain(_llm, _chunks):

    _chunks = chunk_using_nltk(_chunks)

    from langchain.chains import MapReduceDocumentsChain, ReduceDocumentsChain, LLMChain, StuffDocumentsChain
    from langchain.prompts import PromptTemplate

    map_template = """The following is a set of fragments from a product requirements document. 
                      Go through it and summarize it as a list of product requirement summaries. 
                      Use a single sentence for each requirement summary. Provide the response 
                      with each requirement on a different line:
                      ====== begin product requirements document fragments ======
                      {document_fragments}
                      ====== end product requirements document fragments ======
                      List of requirements:
                      """
    map_prompt = PromptTemplate.from_template(map_template)
    map_chain = LLMChain(llm=_llm, prompt=map_prompt)


    reduce_template = """The following contains multiple requirements, with one requirement on each line. 
                         ===== begin requirements  =====
                         {requirements_list}
                         ===== end requirements  =====
                         Take these and identify any duplicate or redundant requirements.
                         Eliminate the redundant and duplicate requirements, and create a new 
                         single list of requirements. Provide the response in a JSON object as a list of strings.
                         Format Instructions:
                         Please provide the output as a JSON using the schema {{ "requirements": [ "requirement1", "requirement2" ] }}
                         List of requirements:
                        """
    from langchain_core.output_parsers import JsonOutputParser
    # parser = JsonOutputParser(pydantic_object=Requirements)
    # print(f"=====\nFormat Instructions: {parser.get_format_instructions()}\n=====")
    reduce_prompt = PromptTemplate(template=reduce_template, 
                                   input_variables=["requirements_list"]
                                #    , partial_variables={"format_instructions": parser.get_format_instructions()}
                                   )
    reduce_chain = LLMChain(llm=_llm, prompt=reduce_prompt)
    combine_documents_chain = StuffDocumentsChain(llm_chain=reduce_chain, document_variable_name="requirements_list")

    reduce_requirements_chain = ReduceDocumentsChain(
        combine_documents_chain=combine_documents_chain,
        collapse_documents_chain=combine_documents_chain,
        token_max = 4000
    )

    map_reduce_chain = MapReduceDocumentsChain(
        llm_chain=map_chain,
        reduce_documents_chain=reduce_requirements_chain,
        document_variable_name="document_fragments",
        return_intermediate_steps=False
    )

    requirements = map_reduce_chain.run(_chunks)

    import json
    # print(requirements)
    # st.subheader("Response as returned")
    # st.write(requirements)
    return json.loads(requirements)["requirements"]


@st.cache_data
def get_story(_chain, req):
    return _chain({"query": req})

if prd_file is not None:
    set_debug(True)
    # File has been uploaded. Parse it
    temp_file_name, document = load_document(prd_file)
    # st.write(f"Total no of tokens for the document: {calculate_tokens_for_document(document)}")

    chunks = parse_word_doc_single_chunk(temp_file_name, document)
    # st.write(f"Length {len(chunks)}")

    # chunks = chunk_using_nltk(chunks)
    # st.write(f"Length {len(chunks)}")

    # show_chunks(calculate_tokens, chunks)

    from langchain_openai import OpenAIEmbeddings
    vectorstore = documents_to_chroma(chunks, OpenAIEmbeddings())

    retriever = vectorstore.as_retriever()

    # requirements = get_requirements_from_stuff_chain(get_gpt35_llm(), chunks)
    # for requirement in requirements:
    #     st.write(requirement)


    # import sqlite3
    # conn = sqlite3.connect('my_database.db')
    # # Create a cursor object
    # c = conn.cursor()

    # # Create a new table
    # # c.execute('''
    # #     CREATE TABLE requirements (
    # #         id INTEGER PRIMARY KEY,
    # #         requirement TEXT
    # #     )
    # # ''')
    requirements = get_requirements_from_mapreduce_chain(get_gpt35_llm(), chunks)
    for i, requirement in enumerate(requirements):
        st.markdown(f"{i+1}. {requirement}")
    #     c.execute('''
    #         INSERT INTO requirements (id, requirement)
    #         VALUES (?, ?)
    #     ''', (i, requirement))

    # conn.commit()
    # conn.close()

    # c.execute('SELECT * FROM requirements')

    # results = c.fetchall()
    # options = [result[1] for result in results]

    # options = 

    selected_reqs = st.multiselect("Select the requirements to expand", requirements, [])

    expand_llm = get_gpt35turbo()

    from langchain.prompts import PromptTemplate
    from langchain.chains import RetrievalQA

    prompt_template = PromptTemplate(
        template="""You are a Product Manager for an Agile program. You create stories for 
                    implementation by the engineering team, that follow the INVEST framework,
                    where the letters stand for
                    'I' ndependent (of all others)
                    'N' egotiable (not a specific contract for features)
                    'V' aluable (or vertical)
                    'E' stimable (to a good approximation)
                    'S' mall (so as to fit within an iteration)
                    'T' estable (in principle, even if there isn’t a test for it yet).
                    Please create a user story for the following requirement:
                    The context for the requirements is given below.
                    ==== begin context ====
                    {context}
                    ==== end context ====
                    ==== begin requirement ====
                    {question}
                    ==== end requirement ====
                    
                    Each story must have a 
                    1. Title
                    2. Description
                    3. List of Acceptance Criteria
                    Provide the output as a JSON using the following schema:
                    {{ "title": "Story Title", "description": "Story Description", "acceptance_criteria": [ "criteria 1", "criteria 2", "criteria 3" ] }}
                    """,
        input_variables=['requirement'])

    # chain = LLMChain(llm=expand_llm, prompt=prompt_template, retriever=retriever)
    chain = RetrievalQA.from_chain_type(
        llm=expand_llm,
        retriever=retriever,
        chain_type_kwargs={"prompt": prompt_template}
    )

    for req in selected_reqs:
        response = get_story(chain, req)
        st.subheader(response["query"])
        import json
        story = json.loads(response["result"])
        st.markdown(f"**Title**: {story['title']}")
        st.markdown(f"**Description**: {story['description']}")
        st.markdown(f"**Acceptance Criteria**")
        for i, criteria in enumerate(story["acceptance_criteria"]):
            st.markdown(f"{i+1}. {criteria}")

    ## Send document to LLM to extract functional requirement stories
